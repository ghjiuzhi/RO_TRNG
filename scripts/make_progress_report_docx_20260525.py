#!/usr/bin/env python3
"""Generate a Chinese DOCX progress report for the RO_TRNG project."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "doc/RO_TRNG_项目进度表_20260525.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BODY_FONT = "等线"
LATIN_FONT = "Calibri"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = LATIN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str = "", style: str | None = None, size: float | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=size)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    if level == 1:
        size, before, after, color = 16, 16, 8, BLUE
    elif level == 2:
        size, before, after, color = 13, 12, 6, BLUE
    else:
        size, before, after, color = 12, 8, 4, DARK_BLUE
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table)
    widths = [1900, 7460]
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        set_cell_width(cells[0], widths[0])
        set_cell_width(cells[1], widths[1])
        set_cell_shading(cells[0], LIGHT_GRAY)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_run_font(run, size=9.5, bold=(cell is cells[0]))


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table)
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_width(hdr_cells[i], widths[i])
        set_cell_shading(hdr_cells[i], LIGHT_BLUE)
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                set_run_font(run, size=8.8, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                if i in (0, 2):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=8.2)


def add_note_box(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.5)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("RO_TRNG 论文实验进度表 | 2026-05-25")
    set_run_font(run, size=8.5, color="666666")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("E:\\Project\\MLDSA\\RO_TRNG")
    set_run_font(run, size=8, color="666666")


def build_doc() -> None:
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("RO_TRNG 论文项目详细进度表")
    set_run_font(r, size=22, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("从原始工程与初步想法，到当前可复现机制证据链")
    set_run_font(r, size=12, color="555555")

    add_kv_table(
        doc,
        [
            ("项目目录", r"E:\Project\MLDSA\RO_TRNG"),
            ("开发板与环境", "Zynq-7020 / 正点原子领航者 v2 复现板；Vivado 2023.2；COM3 @ 115200"),
            ("报告日期", "2026-05-25"),
            ("当前阶段", "已从“能否采到数据”推进到“机制证据链、复现实验入口和论文图表包整理”阶段。"),
            ("当前主线", "placement 敏感性不应简单归因于 RO-RO hard locking；采样端物理实现是熵源边界的一部分。"),
        ],
    )

    add_note_box(
        doc,
        "一句话总结",
        "最开始只有原始工程、复现工程和一个“用 TDC/placement 解释 RO 熵源差异”的想法；现在已经具备真实硬件采集链路、placement/TRNG/RO_FREQ/TDC/SP800-90B restart 数据、sample RO 双向反事实机制证据、可重复生成的论文表格与图表脚本，并形成了明确的高水平投稿缺口清单。",
    )

    add_heading(doc, "1. 起点与当前对比", 1)
    add_table(
        doc,
        ["维度", "最开始状态", "当前状态", "变化意义"],
        [
            [
                "工程基础",
                "只有原始 `fpga/` 工程、原始论文 PDF，以及手动复现的 `fpga1` 工程；串口/JTAG/板级连接状态不稳定。",
                "已明确 Vivado 路径、COM3、UART 采集流程、XADC 读取、bitstream 编程脚本、分析脚本和复现实验入口。",
                "从手动摸索变成可脚本化、可追踪、可复现的实验体系。",
            ],
            [
                "研究问题",
                "初始想法是用 TDC 测 RO 的频率、抖动、相位差，再比较不同 placement 的随机性。",
                "主线收敛为 sampler-side entropy-source boundary：sample RO、采样寄存器、局部路由和采样孔径共同影响熵源。",
                "从“手动 placement 改好坏”提升为机制型论文问题。",
            ],
            [
                "硬件数据",
                "没有稳定的自动采集；Xcom/SScom 数据格式和串口脚本未统一；TDC bitstream 生成来源不清。",
                "已采集 placement matrix、20 MiB repeat、RO_FREQ、pair TDC、reset-aligned TDC、SP800-90B restart 和 sample RO 反事实数据。",
                "已经有真实硬件证据，不再只是仿真或设计设想。",
            ],
            [
                "SP800-90B",
                "原始工程未提供完整可直接复用的 restart 实验工具；最初还不清楚如何生成 1000x1000 restart 矩阵。",
                "已打通 MinGW 版 90B 工具链、auto-stream restart、1000x1000 / 1000x125 bit-symbol restart、warmup 扫描和列偏置分析。",
                "可以把 restart startup transient 写成论文机制证据，而不只是随机性测试结果。",
            ],
            [
                "论文证据",
                "只有“placement 可能影响 RO 熵源”的猜想。",
                "已有强因果证据：compact top + formal sample RO 会拉坏；formal top + compact sample RO 会修好。TDC 同时排除了简单 hard locking。",
                "形成了可以投稿的核心机制闭环。",
            ],
        ],
        [1150, 2680, 3020, 2510],
    )

    add_heading(doc, "2. 阶段进度总表", 1)
    progress_rows = [
        [
            "0",
            "原始材料梳理",
            "阅读原始工程、`fpga1` 复现工程和论文 PDF；明确原始工程并未直接提供完整 90B restart 自动工具。",
            "完成",
            "确定不能把原始论文当作完整实验包，需要自行搭建采集、分析、复现链路。",
        ],
        [
            "1",
            "串口/JTAG 链路打通",
            "修正 ARM DAP 误编程问题、COM3 参数问题、PS/PL 串口连线问题；实现 `capture_uart.ps1`、program/capture 自动化、SHA256 和 metadata。",
            "完成",
            "解决“能不能真实采到板上数据”的第一瓶颈。",
        ],
        [
            "2",
            "placement matrix 与连续流 TRNG",
            "采集 compact、checker、sparse、far、same_column、cross_region、random1/2/3 等 placement；补 10/20 MiB repeats。",
            "完成，仍可扩展",
            "证明 placement 会显著改变连续流随机性，建立 random1 坏例和 random3 好例。",
        ],
        [
            "3",
            "RO_FREQ 与 pair TDC",
            "完成 random1/random3 频率重复、near/far baseline 和 6 个 pair-specific TDC；TDC 相关性接近 0。",
            "完成",
            "TDC 结果弱化“坏 placement=RO-RO 硬锁定”的简单解释。",
        ],
        [
            "4",
            "SP800-90B 工具链",
            "使用 MinGW 构建 `ea_non_iid`、`ea_iid`、`ea_restart`；完成 non-IID smoke、8M bit-symbol 估计和结果汇总。",
            "完成",
            "从简单 p1/STS 走向标准化 entropy assessment。",
        ],
        [
            "5",
            "restart auto-stream",
            "新增 restart auto-stream RTL 和 Vivado in-memory flow；完成 4x64 smoke、1000x1000 byte-symbol、1000x125 bit-symbol formal restart。",
            "完成",
            "把原本需要多天 reprogram 的 restart 数据集降到单次 bitstream 长流采集。",
        ],
        [
            "6",
            "restart warmup 与固定列偏置",
            "完成 random3 warmup0/8/10/11/12/16 与 repeat02；边界复现为 warmup10 fail、warmup11/12 pass。",
            "完成",
            "证明连续流好不等于 restart 初始固定位置稳定；startup transient 成为机制线索。",
        ],
        [
            "7",
            "sampler-side ablation",
            "完成 random1 sampler-island / regs-only 20 MiB confirmation，连续流从强偏置修复到近理想。",
            "完成",
            "首次强烈指向采样端物理实现，而不仅是 data RO 阵列本身。",
        ],
        [
            "8",
            "sample RO 双向反事实",
            "forward：compact top + formal-routed sample RO 变坏；reverse：formal top + compact-routed sample RO 修好。",
            "完成，当前最强",
            "形成因果闭环，是目前最有论文价值的机制证据。",
        ],
        [
            "9",
            "clean reset-aligned TDC",
            "完成 6 点 clean32k TDCR header-aligned TDC：random1 baseline、random3 goodref、random1 sampler-local，各 warmup0/12。",
            "完成",
            "TDC 从不完全对齐的 raw packet 变成可防守的 reset/header-aligned 证据。",
        ],
        [
            "10",
            "论文材料与复现包",
            "生成机制证据链表、clean TDC 图、sample RO 反事实表、claim boundary、图表计划、无多板阶段计划。",
            "进行中",
            "论文写作可以正式开始，后续重点是多板复现与校准补强。",
        ],
    ]
    add_table(doc, ["阶段", "任务", "已完成内容", "状态", "论文意义"], progress_rows, [700, 1350, 3600, 900, 2810])

    add_heading(doc, "3. 已完成的硬件与数据成果", 1)
    add_heading(doc, "3.1 连续流 TRNG 与 placement", 2)
    add_paragraph(
        doc,
        "已完成多类 placement 的真实硬件采集，包括 compact、checker、sparse、far、same_column、cross_region、random1/2/3、row 等，并补充 10 MiB / 20 MiB 级 repeat。关键结论是：在同一板卡、同一结构、同一采集链路下，仅 placement 变化就能显著改变原始随机性。",
    )
    add_table(
        doc,
        ["代表数据", "结果", "意义"],
        [
            ["random1 formal 10 MiB", "p1 = 0.337315512；快速 bit min-entropy = 0.593605945", "稳定坏例，说明 placement 可导致强偏置。"],
            ["random3 formal 10 MiB", "p1 = 0.499968565；快速 bit min-entropy = 0.999909299", "稳定好例，作为 good reference。"],
            ["original fpga1 baseline 10 MiB", "p1 = 0.500035894；快速 bit min-entropy = 0.999896436", "证明复现工程本身采集链路可得到近理想输出。"],
            ["random1/random3 repeat03", "random1 仍低熵；random3 20 MiB p1 = 0.499915", "排除单次采集偶然和 bit-order 假象。"],
        ],
        [2200, 3300, 3860],
    )

    add_heading(doc, "3.2 SP800-90B 与 restart", 2)
    add_paragraph(
        doc,
        "已构建 Windows/MinGW 路线的 SP800-90B 工具链，并完成 non-IID smoke、8M bit-symbol 估计、restart auto-stream 和 formal-size restart 数据集。重要进展是：`1000 x 1000` 和 `1000 x 125 -> 1000 x 1000 bit-symbol` 两条 restart 路线都已经在真实硬件上打通。",
    )
    add_table(
        doc,
        ["项目", "完成情况", "关键结果"],
        [
            ["90B 工具链", "`ea_non_iid.exe`、`ea_iid.exe`、`ea_restart.exe` 均已编译运行", "IID smoke 不适合作为主线；论文应使用 non-IID/restart evidence。"],
            ["8M bit-symbol non-IID", "random1、random3、original baseline 均完成", "random1 H_original 约 0.389520；random3 约 0.902345；original 约 0.877727。"],
            ["auto-stream restart", "4x64 smoke、1000x1000 byte-symbol、1000x125 packed-byte 均完成", "restart fast-path 从设计草图变成真实板级链路。"],
            ["random3 warmup 扫描", "warmup0/8/10 fail；warmup11/12/16 pass；repeat02 复现边界", "startup transient 和固定采样位置偏置成为论文机制证据。"],
        ],
        [2100, 3450, 3810],
    )

    add_heading(doc, "3.3 TDC 与机制诊断", 2)
    add_paragraph(
        doc,
        "TDC 线已经从 near/far 和 pair-specific 探索推进到 reset-aligned / warmup-aligned clean32k 矩阵。当前 TDC 的最稳结论不是“证明锁定”，而是排除简单 pairwise hard locking，并为 sampler-side 机制提供约束。",
    )
    add_table(
        doc,
        ["TDC 项", "完成内容", "结论边界"],
        [
            ["pair-specific TDC", "6 个重点 pair；96 个窗口；strong-lock windows = 0；max small-lag abs correlation 约 0.0318", "没有观察到强 pair-level phase locking；不能说完全无耦合。"],
            ["clean reset-aligned TDC", "6 个 TDCR header-aligned captures，每个 32768 packets", "same-bin ratio 约 1%，longest run = 3，autocorr 接近 0。"],
            ["sampler-local TDC", "random1 sampler-local warmup12 的 H(diff) 和 transition H(diff) 在六点矩阵中最高", "弱正证据；不能单独作为因果证明。"],
        ],
        [2100, 3950, 3310],
    )

    add_heading(doc, "3.4 sample RO 双向反事实", 2)
    add_paragraph(
        doc,
        "这是目前最关键的机制突破：只改变 sample RO 的 routed physical implementation，即可把 restart passband 拉坏或修好。它比普通相关性更强，因为形成了 forward fail 和 reverse repair 的双向反事实闭环。",
    )
    add_table(
        doc,
        ["方向", "top design", "sample RO 实现", "warmup", "overall p1 / min-H", "解释"],
        [
            ["forward fail", "compact FIFO diagnostic", "formal-routed locked", "4", "p1=0.376651；min-H=0.681888", "原本 near-ideal 的 compact passband 被拉回强偏置失败。"],
            ["forward fail", "compact FIFO diagnostic", "formal-routed locked", "5", "p1=0.373430 / 0.373541", "原本通过的 warmup5 也被稳定拉坏，并复现。"],
            ["forward fail", "compact FIFO diagnostic", "formal-routed locked", "11", "p1=0.464819；min-H=0.901901", "长 warmup 后偏置减弱但仍可见。"],
            ["reverse repair", "formal auto restart", "compact-routed locked", "4", "p1=0.499419 / 0.499754；min-H≈0.998-0.999", "formal warmup4 失败被修复到近理想，并复现。"],
        ],
        [1300, 1700, 1800, 800, 1900, 1860],
    )
    add_note_box(
        doc,
        "机制结论",
        "TDC 证据排除了简单 pairwise RO locking 作为主导解释；sample RO 双向反事实显示 sampler-side physical implementation 会重塑 restart warmup passband。因此采样路径不能被视为被动读出电路，而应纳入物理熵源边界。",
    )

    add_heading(doc, "4. 已形成的脚本、文档和可复现产物", 1)
    add_table(
        doc,
        ["类别", "关键文件", "作用"],
        [
            ["采集脚本", "`scripts/capture_uart.ps1`；`scripts/program_and_capture_uart.ps1`；`scripts/program_and_capture_uart_preopen.ps1`", "自动串口采集、Vivado 编程、SHA256、metadata、XADC after-only。"],
            ["90B 工具链", "`scripts/build_90b_mingw.ps1`；`scripts/run_90b_restart.ps1`；`scripts/prepare_90b_inputs.py`", "构建和运行 SP800-90B non-IID/restart 分析。"],
            ["restart RTL", "`rtl/restart/RO_TRNG_restart_auto_top.v`；`RO_TRNG_restart_fifo_compact_diag_top.v`", "正式 restart 长流和 diagnostic/compact 对照。"],
            ["TDC RTL/脚本", "`rtl/tdc/`；`scripts/build_tdc_reset_aligned_bitstreams.ps1`；`scripts/make_tdc_clean32k_figures_20260525.py`", "TDC bitstream 构建、采集和论文图生成。"],
            ["机制证据链", "`data/experiments/mechanism_evidence_chain_20260525/`", "把 sample RO、TDC、XADC 和 restart 证据合并。"],
            ["sample RO 表", "`data/experiments/sample_ro_counterfactual_20260525/`", "论文级 forward fail / reverse repair 汇总表。"],
            ["复现指南", "`doc/reproduce_key_experiments_20260525.md`", "以后重做关键实验的入口文档。"],
            ["论文边界", "`doc/paper_claim_evidence_boundary_20260525.md`；`doc/paper_figure_table_plan_20260525.md`", "明确能写什么、不能写什么，以及正文图表计划。"],
        ],
        [1450, 4610, 3300],
    )

    add_heading(doc, "5. 论文层面的当前成果", 1)
    add_table(
        doc,
        ["论文问题", "现在能支撑的结论", "证据强度"],
        [
            ["placement 是否显著影响 RO-TRNG？", "能。random1/random3/compact/checker/sparse/far 等表现差异明显，并有 repeat。", "强"],
            ["差异是否只是 RO-RO 硬锁定？", "不是简单 hard locking。pair TDC 与 clean TDC 都未显示强持续锁定。", "中强，偏排除性"],
            ["采样端是否是熵源边界？", "是当前最强主张。sample RO 双向反事实能拉坏/修好 restart outcome。", "强"],
            ["连续流好是否足够？", "不够。random3 restart 固定列和 warmup transition 表明 startup/restart 是独立风险。", "强"],
            ["TDC 能否解释机理？", "能约束机理、排除简单模型；未校准前不能写 ps 级绝对 jitter。", "中"],
            ["能否冲高水平？", "已有机制亮点；还需要多板复现、TDC code-density calibration、command-gated before/after XADC。", "有潜力但需补强"],
        ],
        [2300, 5160, 1900],
    )

    add_heading(doc, "6. 目前仍缺的内容", 1)
    add_table(
        doc,
        ["缺口", "为什么重要", "建议优先级", "下一步"],
        [
            ["多板复现", "回应 single-board anecdote 质疑，是高水平投稿前最关键外部验证。", "P0", "board B/C 跑 sample RO forward fail 和 reverse repair，每项至少 2 次。"],
            ["TDC code-density calibration", "没有它不能写绝对 ps 级 phase/jitter；当前只能 raw-bin 相对比较。", "P1", "新增独立 calibration top，8-16 MiB per mode。"],
            ["command-gated capture", "解决 auto-stream 与 XADC before_capture/JTAG 延迟冲突。", "P1", "先确认 fpga1 UART_RX 引脚，再做 echo smoke 和 restart/TDC command-gated top。"],
            ["完整论文初稿", "当前数据已经足够开始写，不应等全部硬件完成再写。", "P0", "先写 Results/Discussion，明确边界和 limitation。"],
            ["GitHub/export 版本管理", "保证后续 GPT/Claude/导师分析时看到可复现脚本和表格，不上传巨大 raw bin。", "P1", "在重要阶段打 Git snapshot 或更新 export。"],
        ],
        [1800, 3300, 900, 3360],
    )

    add_heading(doc, "7. 推荐下一步时间表", 1)
    add_table(
        doc,
        ["时间段", "目标", "具体动作", "交付物"],
        [
            ["现在，无多板", "把单板结果变成论文材料", "写 Results/Discussion 初稿；整理图表；完善 claim boundary。", "论文初稿 v1；正文图表包。"],
            ["下一次上板", "提高实验严谨性", "做 command-gated UART RX smoke；必要时补 XADC before/after；避免再依赖长 start delay。", "command-gated smoke report。"],
            ["多板可用后", "验证核心机制是否跨板", "优先跑 sample RO forward fail / reverse repair，而不是全 placement 乱跑。", "multi-board counterfactual table。"],
            ["投稿前补强", "让 TDC 结论更硬", "完成 dedicated code-density calibration，重新分析 clean TDC/pair TDC。", "calibrated TDC figure/table。"],
        ],
        [1500, 2250, 3710, 1900],
    )

    add_heading(doc, "8. 当前结论", 1)
    add_paragraph(
        doc,
        "这个项目已经从“复现一个 RO_TRNG 工程并尝试手动 placement”推进到“围绕 sampler-side entropy-source boundary 的机制验证”。目前最有价值的结论不是手动 placement 本身，而是：采样端物理实现会改变 restart passband 和连续流随机性；TDC 没有支持简单硬锁定解释，反而帮助把机制范围收缩到 sample RO、采样寄存器、局部路由和采样孔径。",
    )
    add_paragraph(
        doc,
        "以当前结果，已经可以开始写一篇有机制亮点的会议或期刊初稿。若要冲高水平期刊，需要把 sample RO 双向反事实扩展到多板，并补独立 TDC code-density calibration 与 command-gated capture。",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(f"Wrote {OUT}")
